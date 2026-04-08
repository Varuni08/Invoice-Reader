import re
import json
import pandas as pd
import pdfplumber
import PyPDF2
from docx import Document
from openpyxl import load_workbook
from io import BytesIO


INVOICE_START_PATTERN = re.compile(
    r'(invoice\s*#|invoice\s*no|invoice\s*number|bill\s*to|remit\s*to)',
    re.IGNORECASE
)


def split_pdf_into_invoices(pages: list) -> list:
    if not pages:
        return []
    chunks = []
    current = [pages[0]]
    for page in pages[1:]:
        if INVOICE_START_PATTERN.search(page):
            chunks.append("\n\n".join(current))
            current = [page]
        else:
            current.append(page)
    chunks.append("\n\n".join(current))
    return chunks


def parse_pdf(file_bytes: bytes) -> list:
    pages = []
    try:
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                pages.append(text)
    except Exception:
        reader = PyPDF2.PdfReader(BytesIO(file_bytes))
        for page in reader.pages:
            pages.append(page.extract_text() or "")
    return split_pdf_into_invoices(pages)


def parse_xlsx(file_bytes: bytes) -> list:
    wb = load_workbook(BytesIO(file_bytes), read_only=True, data_only=True)
    chunks = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            row_vals = [str(c) if c is not None else "" for c in row]
            if any(v.strip() for v in row_vals):
                rows.append("\t".join(row_vals))
        if rows:
            chunks.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows))
    return chunks if chunks else [""]


def parse_docx(file_bytes: bytes) -> list:
    doc = Document(BytesIO(file_bytes))
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("\t".join(cells))
    return ["\n".join(lines)]


def parse_csv(file_bytes: bytes) -> list:
    try:
        df = pd.read_csv(BytesIO(file_bytes))
        return [df.to_string(index=False)]
    except Exception:
        return [file_bytes.decode("utf-8", errors="replace")]


def parse_txt(file_bytes: bytes) -> list:
    return [file_bytes.decode("utf-8", errors="replace")]


def parse_json(file_bytes: bytes) -> list:
    try:
        data = json.loads(file_bytes.decode("utf-8", errors="replace"))
        return [json.dumps(data, indent=2)]
    except Exception:
        return [file_bytes.decode("utf-8", errors="replace")]


def parse_file(filename: str, file_bytes: bytes) -> list:
    ext = filename.lower().rsplit(".", 1)[-1]
    parsers = {
        "pdf":  parse_pdf,
        "xlsx": parse_xlsx,
        "xls":  parse_xlsx,
        "docx": parse_docx,
        "doc":  parse_docx,
        "csv":  parse_csv,
        "tsv":  parse_csv,
        "txt":  parse_txt,
        "json": parse_json,
    }
    fn = parsers.get(ext, parse_txt)
    chunks = fn(file_bytes)
    return [c.strip() for c in chunks if c.strip()]
